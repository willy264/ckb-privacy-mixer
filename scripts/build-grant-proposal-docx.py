#!/usr/bin/env python3
"""Build a black-and-white styled proposal DOCX with original-color media."""

from __future__ import annotations

import argparse
import hashlib
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, Sequence
from urllib.parse import urlparse
from xml.etree import ElementTree

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "grant-proposal.md"
OUTPUT = ROOT / "Obscell_CKB_Community_DAO_Proposal_15K_Grant_Ready.docx"
DIAGRAM_DIR = ROOT / "docs" / "diagrams"
REPOSITORY_BLOB_BASE = "https://github.com/willy264/ckb-privacy-mixer/blob/main/"

BLACK = "000000"
INK = "10241F"
GREEN = "176A4A"
GREEN_DARK = "15483A"
GREEN_LIGHT = "E8F3EE"
PURPLE = "654AB6"
PURPLE_LIGHT = "F0ECFA"
GOLD = "B97900"
GOLD_LIGHT = "FFF5DC"
GRAY = "52615B"
GRAY_LIGHT = "F2F5F3"
WHITE = "FFFFFF"

EXPECTED_HEADINGS = [
    "1. Title",
    "2. Summary",
    "3. Project Introduction",
    "4. Team & Roles",
    "5. Current Status",
    "6. Application Design",
    "7. Key Benefits for CKB",
    "8. Detailed Deliverables & Milestones",
    "9. Budget Breakdown",
    "10. Out-of-Scope / Future Funding Needs",
    "11. Risk & Mitigation",
    "12. Closing / Call to Action",
    "13. Supporting Links",
]

PAGE_BREAK_HEADINGS = {
    "3. Project Introduction",
    "5. Current Status",
    "6. Application Design",
    "8. Detailed Deliverables & Milestones",
    "9. Budget Breakdown",
    "11. Risk & Mitigation",
    "13. Supporting Links",
}

DIAGRAM_FILES = {
    "system": DIAGRAM_DIR / "system-architecture.png",
    "deposit": DIAGRAM_DIR / "deposit-flow.png",
    "withdrawal": DIAGRAM_DIR / "withdrawal-flow.png",
    "state": DIAGRAM_DIR / "state-vault-relationship.png",
    "sdk": DIAGRAM_DIR / "sdk-integration.png",
    "trust": DIAGRAM_DIR / "trust-boundary.png",
}

EVIDENCE_FILES = [
    ROOT / "docs" / "evidence" / f"figure-{number}-{suffix}.png"
    for number, suffix in (
        (1, "legacy-mixer"),
        (2, "ccc-demo"),
        (3, "private-balance"),
        (4, "developer-protocol"),
        (6, "second-consumer"),
    )
]


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    name = "segoeuib.ttf" if bold else "segoeui.ttf"
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default(size=size)


def wrap_pixels(draw: ImageDraw.ImageDraw, value: str, face: ImageFont.ImageFont, width: int) -> list[str]:
    words = value.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if not current or draw.textbbox((0, 0), candidate, font=face)[2] <= width:
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_wrapped_center(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    value: str,
    face: ImageFont.ImageFont,
    fill: str,
    gap: int = 8,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_pixels(draw, value, face, x2 - x1 - 38)
    heights = [draw.textbbox((0, 0), line, font=face)[3] for line in lines]
    total = sum(heights) + gap * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total) / 2
    for line, line_height in zip(lines, heights):
        width = draw.textbbox((0, 0), line, font=face)[2]
        draw.text((x1 + (x2 - x1 - width) / 2, y), line, font=face, fill=fill)
        y += line_height + gap


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    body: str = "",
    fill: str = "FFFFFF",
    outline: str = GREEN,
) -> None:
    draw.rounded_rectangle(box, radius=20, fill=f"#{fill}", outline=f"#{outline}", width=4)
    x1, y1, x2, y2 = box
    compact = x2 - x1 < 400
    title_face = font(28 if compact else 32, bold=True)
    body_face = font(20 if compact else 23)
    if body:
        title_lines = wrap_pixels(draw, title, title_face, x2 - x1 - 38)
        body_lines = wrap_pixels(draw, body, body_face, x2 - x1 - 44)
        title_heights = [draw.textbbox((0, 0), line, font=title_face)[3] for line in title_lines]
        body_heights = [draw.textbbox((0, 0), line, font=body_face)[3] for line in body_lines]
        title_total = sum(title_heights) + 7 * max(0, len(title_lines) - 1)
        body_total = sum(body_heights) + 5 * max(0, len(body_lines) - 1)
        content_height = title_total + 20 + body_total
        y = y1 + max(14, (y2 - y1 - content_height) / 2)
        for line, line_height in zip(title_lines, title_heights):
            line_width = draw.textbbox((0, 0), line, font=title_face)[2]
            draw.text((x1 + (x2 - x1 - line_width) / 2, y), line, font=title_face, fill=f"#{INK}")
            y += line_height + 7
        y += 13
        for line, line_height in zip(body_lines, body_heights):
            line_width = draw.textbbox((0, 0), line, font=body_face)[2]
            draw.text((x1 + (x2 - x1 - line_width) / 2, y), line, font=body_face, fill=f"#{GRAY}")
            y += line_height + 5
    else:
        draw_wrapped_center(draw, box, title, title_face, f"#{INK}")


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str = GREEN,
    width: int = 7,
) -> None:
    draw.line((start, end), fill=f"#{color}", width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - direction * 22, y2 - 14), (x2 - direction * 22, y2 + 14)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 14, y2 - direction * 22), (x2 + 14, y2 - direction * 22)]
    draw.polygon(points, fill=f"#{color}")


def canvas(height: int, title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1800, height), "#FBFCFB")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1800, 66), fill=f"#{INK}")
    draw.text((48, 17), title, font=font(29, bold=True), fill="#FFFFFF")
    label = "TARGET ARCHITECTURE - NOT DEPLOYMENT EVIDENCE"
    label_width = draw.textbbox((0, 0), label, font=font(20, bold=True))[2]
    draw.rounded_rectangle((1740 - label_width, 13, 1770, 52), radius=12, fill=f"#{GOLD}")
    draw.text((1755 - label_width, 20), label, font=font(20, bold=True), fill="#FFFFFF")
    return image, draw


def save_diagram(image: Image.Image, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, format="PNG", optimize=True)


def build_system_diagram() -> None:
    image, draw = canvas(1310, "System architecture")
    boxes = [
        ((430, 115, 1370, 280), "Existing CKB application", "UI, business logic, wallet connector, CCC Client and Signer", WHITE, GREEN),
        ((430, 345, 1370, 510), "Obscell PrivacyClient", "Notes, prover, verified state, protocol planning and typed operations", GREEN_LIGHT, GREEN),
        ((430, 575, 1370, 740), "Corrected fixed-denomination V1", "PoolState, Vault, Staging/refund, nullifier, proof and CT rules", PURPLE_LIGHT, PURPLE),
        ((430, 805, 1370, 970), "CKB consensus", "Canonical cells, transaction conflicts, script validation and confirmation", WHITE, GREEN_DARK),
        ((430, 1035, 1370, 1200), "CKB Pudge", "Grant validation network and independently reviewable evidence", GOLD_LIGHT, GOLD),
    ]
    for box, title, body, fill, outline in boxes:
        rounded_box(draw, box, title, body, fill, outline)
    for first, second in zip(boxes, boxes[1:]):
        arrow(draw, ((first[0][0] + first[0][2]) // 2, first[0][3]), ((second[0][0] + second[0][2]) // 2, second[0][1]))
    rounded_box(draw, (40, 555, 340, 735), "Coordinator", "Derives acceptance candidates from chain", GRAY_LIGHT, GRAY)
    rounded_box(draw, (1460, 555, 1760, 735), "Relayer / indexer", "Reconstructs, submits, observes and caches", GRAY_LIGHT, GRAY)
    arrow(draw, (340, 645), (430, 645), GRAY, 5)
    arrow(draw, (1460, 645), (1370, 645), GRAY, 5)
    save_diagram(image, DIAGRAM_FILES["system"])


def build_horizontal_flow(path: Path, title: str, steps: Sequence[tuple[str, str, str, str]]) -> None:
    image, draw = canvas(600, title)
    margin = 35
    gap = 34
    count = len(steps)
    width = int((1800 - 2 * margin - gap * (count - 1)) / count)
    y1, y2 = 150, 500
    for index, (heading, body, fill, outline) in enumerate(steps):
        x1 = margin + index * (width + gap)
        x2 = x1 + width
        rounded_box(draw, (x1, y1, x2, y2), heading, body, fill, outline)
        if index < count - 1:
            arrow(draw, (x2 + 4, (y1 + y2) // 2), (x2 + gap - 4, (y1 + y2) // 2))
    save_diagram(image, path)


def build_state_diagram() -> None:
    image, draw = canvas(930, "PoolState and Vault provenance")
    rounded_box(draw, (100, 130, 760, 335), "PoolState input", "Sequence n, accepted root R, nullifier root N, immutable pool configuration", PURPLE_LIGHT, PURPLE)
    rounded_box(draw, (1040, 130, 1700, 335), "Sibling Vault input", "Exact CT type A, balance/accounting Q, pinned PoolState identity", GREEN_LIGHT, GREEN)
    rounded_box(draw, (590, 400, 1210, 585), "Atomic transition", "Initialization, acceptance or withdrawal must validate the entire pair", GOLD_LIGHT, GOLD)
    rounded_box(draw, (100, 660, 760, 860), "PoolState output", "Sequence n + 1 and validated root/nullifier/accounting successor", PURPLE_LIGHT, PURPLE)
    rounded_box(draw, (1040, 660, 1700, 860), "Sibling Vault output", "Same CT identity and exact validated balance/conservation successor", GREEN_LIGHT, GREEN)
    arrow(draw, (430, 335), (730, 400))
    arrow(draw, (1370, 335), (1070, 400))
    arrow(draw, (730, 585), (430, 660))
    arrow(draw, (1070, 585), (1370, 660))
    save_diagram(image, DIAGRAM_FILES["state"])


def build_sdk_diagram() -> None:
    image, draw = canvas(1040, "SDK integration boundary")
    rounded_box(draw, (80, 140, 580, 340), "Host application", "UI, connector, public transactions and product policy", WHITE, GREEN)
    rounded_box(draw, (650, 140, 1150, 340), "createPrivacyClient", "Explicit opt-in using injected dependencies", GREEN_LIGHT, GREEN)
    rounded_box(draw, (1220, 140, 1720, 340), "Privacy core", "Protocol, notes, prover, state and services", PURPLE_LIGHT, PURPLE)
    arrow(draw, (580, 240), (650, 240))
    arrow(draw, (1150, 240), (1220, 240))
    rounded_box(draw, (650, 470, 1150, 650), "CCC adapter", "Deployment, reader, transaction, signer and capacity", GOLD_LIGHT, GOLD)
    arrow(draw, (1470, 340), (1050, 470))
    rounded_box(draw, (250, 770, 750, 950), "Application-owned CCC Client", "Network and chain access remain with the application", WHITE, GREEN_DARK)
    rounded_box(draw, (1050, 770, 1550, 950), "Operation-scoped CCC Signer", "User approval remains with the selected wallet", WHITE, GREEN_DARK)
    arrow(draw, (830, 650), (500, 770))
    arrow(draw, (970, 650), (1300, 770))
    save_diagram(image, DIAGRAM_FILES["sdk"])


def build_trust_diagram() -> None:
    image, draw = canvas(1130, "Trust boundary")
    columns = [
        (45, 570, "User-controlled local", GREEN_LIGHT, GREEN, ["Note/nullifier secrets", "Encrypted note store", "Local prover", "Recipient choice", "Wallet signer"]),
        (615, 1185, "Operational / untrusted", GRAY_LIGHT, GRAY, ["Coordinator", "Relayer", "Indexer and HTTP", "Redis queues/cache", "Candidate transactions"]),
        (1230, 1755, "CKB authority", PURPLE_LIGHT, PURPLE, ["Canonical PoolState", "Canonical Vault", "Accepted roots", "Nullifier state", "Verifier and CT scripts"]),
    ]
    for x1, x2, heading, fill, outline, items in columns:
        draw.rounded_rectangle((x1, 125, x2, 990), radius=24, fill=f"#{fill}", outline=f"#{outline}", width=5)
        draw_wrapped_center(draw, (x1 + 15, 145, x2 - 15, 235), heading, font(34, bold=True), f"#{INK}")
        y = 280
        for item in items:
            draw.rounded_rectangle((x1 + 45, y, x2 - 45, y + 105), radius=14, fill="#FFFFFF", outline=f"#{outline}", width=3)
            draw_wrapped_center(draw, (x1 + 55, y + 5, x2 - 55, y + 100), item, font(25), f"#{INK}")
            y += 130
    arrow(draw, (570, 555), (615, 555), GRAY, 5)
    arrow(draw, (1185, 555), (1230, 555), GRAY, 5)
    draw.text((52, 1030), "Public commitments and typed intents may cross rightward. Private secrets never do.", font=font(25, bold=True), fill=f"#{INK}")
    save_diagram(image, DIAGRAM_FILES["trust"])


def build_diagrams() -> None:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    build_system_diagram()
    build_horizontal_flow(
        DIAGRAM_FILES["deposit"],
        "Deposit and acceptance flow",
        [
            ("User-owned 100 CT", "A pre-existing supported CT cell", WHITE, GREEN),
            ("CCC-signed staging", "User authorizes commitment, refund and pool context", GREEN_LIGHT, GREEN),
            ("Canonical confirmation", "Coordinator observes the staging output from chain", GRAY_LIGHT, GRAY),
            ("Atomic acceptance", "Live PoolState and Vault consume confirmed staging", PURPLE_LIGHT, PURPLE),
            ("Accepted note", "Commitment enters the authoritative root; Vault gains 100 CT", GOLD_LIGHT, GOLD),
        ],
    )
    build_horizontal_flow(
        DIAGRAM_FILES["withdrawal"],
        "Withdrawal and recipient-spend flow",
        [
            ("Accepted note", "User retains secrets in encrypted local state", GREEN_LIGHT, GREEN),
            ("Local proof", "Membership, nullifier, recipient and action are bound", PURPLE_LIGHT, PURPLE),
            ("Typed intent", "Direct CCC submission or fee-only relayer", GRAY_LIGHT, GRAY),
            ("Atomic transition", "Nullifier is spent and Vault decreases by exactly 100 CT", PURPLE_LIGHT, PURPLE),
            ("Recipient 100 CT", "Exact recipient-controlled output", GOLD_LIGHT, GOLD),
            ("CCC subsequent spend", "Independent recipient signer proves output usability", WHITE, GREEN_DARK),
        ],
    )
    build_state_diagram()
    build_sdk_diagram()
    build_trust_diagram()


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color: str = BLACK, size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set(qn("xml:space"), "preserve")
    instruction_node.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instruction_node, separate, text_node, end])


def add_hyperlink(paragraph, label: str, target: str) -> None:
    relation_id = paragraph.part.relate_to(
        target,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLACK)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    run.append(run_properties)
    value = OxmlElement("w:t")
    value.text = label
    run.append(value)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|\*[^*]+\*)")


def portable_link(target: str) -> str:
    parsed = urlparse(target)
    if parsed.scheme in {"http", "https", "mailto"}:
        return target
    resolved = (SOURCE.parent / target).resolve()
    try:
        relative = resolved.relative_to(ROOT).as_posix()
    except ValueError:
        return target
    return REPOSITORY_BLOB_BASE + relative


def add_inline(paragraph, value: str, base_bold: bool = False, base_italic: bool = False) -> None:
    position = 0
    for match in INLINE_PATTERN.finditer(value):
        if match.start() > position:
            run = paragraph.add_run(value[position : match.start()])
            run.bold = base_bold
            run.italic = base_italic
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.italic = base_italic
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = rgb(BLACK)
            run.font.size = Pt(9)
        elif token.startswith("["):
            link_match = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                add_hyperlink(paragraph, link_match.group(1), portable_link(link_match.group(2)))
        else:
            run = paragraph.add_run(token[1:-1])
            run.bold = base_bold
            run.italic = True
        position = match.end()
    if position < len(value):
        run = paragraph.add_run(value[position:])
        run.bold = base_bold
        run.italic = base_italic


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(17)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.08

    heading_sizes = {1: 19, 2: 14, 3: 11.5}
    for level, size in heading_sizes.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(BLACK)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(6)

    if "Caption" not in styles:
        styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption = styles["Caption"]
    caption.font.name = "Aptos"
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = rgb(BLACK)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True
    caption.paragraph_format.space_after = Pt(10)

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Aptos"
        style.font.size = Pt(9.4)
        style.font.color.rgb = rgb(BLACK)
        style.paragraph_format.space_after = Pt(3)

    for style_name in ("Hyperlink", "FollowedHyperlink"):
        try:
            style = styles[style_name]
        except KeyError:
            style = styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
        style.font.color.rgb = rgb(BLACK)
        style.font.underline = True

    header = section.header.paragraphs[0]
    header.text = "OBSCELL  /  CKB COMMUNITY DAO PROPOSAL"
    header.style = normal
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(7.5)
        run.font.bold = True
        run.font.color.rgb = rgb(BLACK)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = footer.add_run("Obscell  |  September 2026  |  ")
    prefix.font.size = Pt(7.5)
    prefix.font.color.rgb = rgb(BLACK)
    add_field(footer, "PAGE", "1")
    footer.add_run(" of ")
    add_field(footer, "NUMPAGES", "1")

    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    document.settings.element.append(update_fields)


def add_cover(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(36)
    run = paragraph.add_run("OBSCELL")
    run.font.name = "Aptos Display"
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = rgb(BLACK)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(10)
    title_run = title.add_run("Reusable Privacy Infrastructure\nfor CKB Applications")
    title_run.font.name = "Aptos Display"
    title_run.font.size = Pt(30)
    title_run.font.bold = True
    title_run.font.color.rgb = rgb(BLACK)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(24)
    subtitle_run = subtitle.add_run("CKB Community DAO Proposal  |  Application / Project Development")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = rgb(BLACK)

    table = document.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(5.2)
    table.columns[1].width = Cm(11.2)
    cover_rows = [
        ("DISCUSSION TITLE", "[DIS] Obscell - Reusable Privacy Infrastructure for CKB Applications"),
        ("FUNDING REQUEST", "$15,000 USD equivalent"),
        ("DELIVERY WINDOW", "4 months / 16 weeks"),
        ("VALIDATION NETWORK", "CKB Pudge testnet"),
        ("PRIMARY DELIVERABLE", "Obscell Privacy SDK + corrected fixed-denomination V1"),
    ]
    for row, values in zip(table.rows, cover_rows):
        prevent_row_split(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.width = table.columns[index].width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 130, 140, 130, 140)
            set_cell_shading(cell, WHITE)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline(p, value, base_bold=index == 0)
            for run in p.runs:
                run.font.size = Pt(8.5 if index == 0 else 10)
                if index == 0:
                    run.font.color.rgb = rgb(BLACK)
    set_table_borders(table)

    callout = document.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.LEFT
    callout.autofit = False
    cell = callout.cell(0, 0)
    cell.width = Cm(16.4)
    set_cell_shading(cell, WHITE)
    set_cell_margins(cell, 180, 180, 180, 180)
    set_table_borders(callout, BLACK, "10")
    p = cell.paragraphs[0]
    add_inline(
        p,
        "PROPOSAL, NOT COMPLETION REPORT. The current implementation is a tested, fail-closed foundation. Corrected-V1 Pudge deployment, live settlement, recipient spend, and independent review remain grant deliverables.",
        base_bold=True,
    )
    for run in p.runs:
        run.font.size = Pt(9.4)
        run.font.color.rgb = rgb(BLACK)

    date = document.add_paragraph()
    date.paragraph_format.space_before = Pt(30)
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date.add_run("Prepared September 2026\nWilliams Oluwagbemi Akinwamide  |  willy264")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = rgb(BLACK)
    document.add_page_break()

    toc_heading = document.add_paragraph()
    toc_heading.paragraph_format.space_after = Pt(12)
    toc_run = toc_heading.add_run("Contents")
    toc_run.font.name = "Aptos Display"
    toc_run.font.size = Pt(22)
    toc_run.font.bold = True
    toc_run.font.color.rgb = rgb(BLACK)
    toc = document.add_paragraph()
    add_field(toc, 'TOC \\o "1-2" \\h \\z \\u', "Right-click and update field to populate the table of contents.")
    document.add_page_break()


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator_row(cells: Sequence[str]) -> bool:
    return all(bool(re.fullmatch(r":?-{3,}:?", cell.replace(" ", ""))) for cell in cells)


def add_table(document: Document, rows: Sequence[Sequence[str]]) -> None:
    if not rows:
        return
    column_count = len(rows[0])
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    font_size = 8.0 if column_count >= 5 else 8.2 if column_count == 4 else 8.5 if column_count == 3 else 9.0
    header = rows[0]
    column_widths: list[float] | None = None
    if header and header[0] == "Area" and column_count == 4:
        column_widths = [3.1, 2.5, 5.2, 6.2]
    elif header and header[0] == "Milestone" and column_count == 5:
        column_widths = [2.7, 1.8, 5.3, 5.3, 2.0]
    elif header and header[0] == "Workstream" and column_count == 6:
        column_widths = [6.0, 2.1, 2.1, 2.1, 2.1, 2.4]
    if column_widths:
        table.autofit = False
        for column, width in zip(table.columns, column_widths):
            column.width = Cm(width)
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        prevent_row_split(row)
        if row_index == 0:
            set_repeat_table_header(row)
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if column_widths:
                cell.width = Cm(column_widths[column_index])
                set_cell_margins(cell, 70, 75, 70, 75)
            else:
                set_cell_margins(cell)
            set_cell_shading(cell, WHITE)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, value, base_bold=row_index == 0)
            for run in paragraph.runs:
                run.font.size = Pt(font_size)
                run.font.color.rgb = rgb(BLACK)
    table.rows[-1]._tr.addnext(OxmlElement("w:bookmarkEnd")) if False else None
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(document: Document, value: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, WHITE)
    set_cell_margins(cell, 150, 180, 150, 180)
    set_table_borders(table, BLACK, "9")
    paragraph = cell.paragraphs[0]
    add_inline(paragraph, value)
    for run in paragraph.runs:
        run.font.size = Pt(9.4)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(document: Document, lines: Sequence[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, WHITE)
    set_cell_margins(cell, 130, 160, 130, 160)
    set_table_borders(table, BLACK, "6")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8.2)
    run.font.color.rgb = rgb(BLACK)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_picture(document: Document, path: Path, alt_text: str) -> None:
    with Image.open(path) as source_image:
        width_px, height_px = source_image.size
    max_width_mm = 166
    max_height_mm = 205
    width_mm = max_width_mm
    height_mm = width_mm * height_px / width_px
    if height_mm > max_height_mm:
        height_mm = max_height_mm
        width_mm = height_mm * width_px / height_px
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Mm(width_mm), height=Mm(height_mm))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", alt_text)


def diagram_for_mermaid(lines: Sequence[str]) -> tuple[Path, str]:
    value = "\n".join(lines)
    if "Existing CKB application" in value:
        return DIAGRAM_FILES["system"], "Target Obscell system architecture"
    if "User-owned 100 CT" in value:
        return DIAGRAM_FILES["deposit"], "Target deposit and acceptance flow"
    if "Accepted local note" in value:
        return DIAGRAM_FILES["withdrawal"], "Target withdrawal and recipient-spend flow"
    if "PoolState input" in value:
        return DIAGRAM_FILES["state"], "Target PoolState and Vault provenance relationship"
    if "Host CKB application" in value:
        return DIAGRAM_FILES["sdk"], "Target PrivacyClient and CCC integration boundary"
    if "User-controlled local boundary" in value:
        return DIAGRAM_FILES["trust"], "Target protocol trust boundary"
    raise ValueError(f"No rendered diagram mapping for Mermaid block:\n{value}")


def add_heading(document: Document, value: str, level: int) -> None:
    paragraph = document.add_heading(level=level)
    add_inline(paragraph, value)
    if level == 1 and value in PAGE_BREAK_HEADINGS:
        paragraph.paragraph_format.page_break_before = True


def add_body_paragraph(document: Document, value: str) -> None:
    caption_match = re.fullmatch(r"\*([^*]*(?:Figure|diagram)[^*]*)\*", value, flags=re.IGNORECASE)
    if caption_match:
        paragraph = document.add_paragraph(style="Caption")
        add_inline(paragraph, caption_match.group(1), base_italic=True)
        return
    paragraph = document.add_paragraph()
    add_inline(paragraph, value)


def render_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            add_heading(document, heading.group(2), len(heading.group(1)))
            index += 1
            continue
        if stripped.startswith("```"):
            language = stripped[3:].strip().lower()
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            index += 1
            if language == "mermaid":
                diagram, alt = diagram_for_mermaid(code_lines)
                add_picture(document, diagram, alt)
            else:
                add_code(document, code_lines)
            continue
        if stripped.startswith("|") and index + 1 < len(lines):
            first = parse_table_row(stripped)
            separator = parse_table_row(lines[index + 1])
            if len(first) == len(separator) and is_separator_row(separator):
                table_rows: list[list[str]] = [first]
                index += 2
                while index < len(lines) and lines[index].strip().startswith("|"):
                    row = parse_table_row(lines[index])
                    if len(row) != len(first):
                        raise ValueError(f"Malformed Markdown table at line {index + 1}")
                    table_rows.append(row)
                    index += 1
                add_table(document, table_rows)
                continue
        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            image_path = (SOURCE.parent / image_match.group(2)).resolve()
            if not image_path.exists():
                raise FileNotFoundError(image_path)
            add_picture(document, image_path, image_match.group(1))
            index += 1
            continue
        if stripped.startswith(">"):
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip()[1:].strip())
                index += 1
            add_callout(document, " ".join(quote_lines))
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, numbered.group(1))
            index += 1
            continue
        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                break
            if re.match(r"^(#{1,3})\s+", candidate) or candidate.startswith(("```", "|", ">", "- ", "![")) or re.match(r"^\d+\.\s+", candidate):
                break
            paragraph_lines.append(candidate)
            index += 1
        add_body_paragraph(document, " ".join(paragraph_lines))


def validate_source(markdown: str) -> None:
    headings = re.findall(r"^#\s+(.+)$", markdown, flags=re.MULTILINE)
    if headings != EXPECTED_HEADINGS:
        raise ValueError(f"Expected exactly the 13 DAO headings in order; found: {headings}")
    if markdown.count("[TO BE PROVIDED BEFORE SUBMISSION]") != 1:
        raise ValueError("The one unresolved funding-address field must remain explicit and unique.")
    if "figure-5-pudge-e2e.png" in markdown.lower():
        raise ValueError("Figure 5 media must remain absent until verified Pudge E2E evidence exists.")
    required_images = [f"evidence/figure-{number}" for number in (1, 2, 3, 4, 6)]
    for marker in required_images:
        if marker not in markdown:
            raise ValueError(f"Missing required evidence reference: {marker}")
    for stale_claim in ("multiple supported assets", "multiple configured pools/assets/values", "fully private", "production-ready"):
        if stale_claim.lower() in markdown.lower():
            raise ValueError(f"Stale or unsupported claim remains: {stale_claim}")


def validate_docx(path: Path) -> None:
    document = Document(path)
    headings = [p.text for p in document.paragraphs if p.style.name == "Heading 1"]
    if headings != EXPECTED_HEADINGS:
        raise ValueError(f"DOCX heading drift: {headings}")
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                values.extend(paragraph.text for paragraph in cell.paragraphs)
    text = "\n".join(values)
    if "Figure 5 - Corrected-V1 Pudge E2E: intentionally absent" not in text:
        raise ValueError("DOCX lost the Figure 5 absence disclosure.")
    if text.count("[TO BE PROVIDED BEFORE SUBMISSION]") != 1:
        raise ValueError("DOCX funding-address placeholder count changed.")
    relationships = list(document.part.rels.values())
    image_relations = [rel for rel in relationships if rel.reltype.endswith("/image")]
    hyperlink_relations = [rel for rel in relationships if rel.reltype.endswith("/hyperlink")]
    if len(image_relations) != 11:
        raise ValueError(f"Expected 11 embedded images (6 diagrams + 5 evidence figures), found {len(image_relations)}")
    source_images = [*DIAGRAM_FILES.values(), *EVIDENCE_FILES]
    expected_hashes = sorted(hashlib.sha256(source.read_bytes()).digest() for source in source_images)
    embedded_hashes = sorted(hashlib.sha256(rel.target_part.blob).digest() for rel in image_relations)
    if embedded_hashes != expected_hashes:
        raise ValueError("Embedded proposal images must remain byte-identical to their original color sources.")
    if len(hyperlink_relations) < 10:
        raise ValueError(f"Expected real hyperlinks, found only {len(hyperlink_relations)}")
    for table in document.tables:
        borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        if borders is None:
            raise ValueError("Every proposal table must define explicit borders.")
        for border in borders:
            if border.get(qn("w:color"), "").upper() != BLACK:
                raise ValueError("Every proposal table border must be black.")
        for row in table.rows:
            for cell in row.cells:
                shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
                if shading is None or shading.get(qn("w:fill"), "").upper() != WHITE:
                    raise ValueError("Every proposal table cell must have a white background.")
    document_xml = document.part.element.xml
    visible_xml = [document_xml]
    for section in document.sections:
        visible_xml.extend((section.header._element.xml, section.footer._element.xml))
    for xml in visible_xml:
        for color_tag in re.findall(r"<w:color\b[^>]*/>", xml):
            value = re.search(r'w:val="([^"]+)"', color_tag)
            if value is not None and value.group(1).upper() not in {BLACK, "AUTO"}:
                raise ValueError(f"Visible proposal text contains a non-black color: {value.group(1)}")
    for field in ("TOC", "PAGE", "NUMPAGES"):
        combined_xml = document_xml + "".join(section.footer._element.xml for section in document.sections)
        if field not in combined_xml:
            raise ValueError(f"Missing Word field: {field}")
    with zipfile.ZipFile(path) as package:
        settings_xml = package.read("word/settings.xml")
        styles_xml = package.read("word/styles.xml").decode("utf-8")
    hyperlink_styles = 0
    for style_id in ("Hyperlink", "FollowedHyperlink", "Hyperlink1", "FollowedHyperlink1"):
        style_blocks = re.findall(
            rf'<w:style\b[^>]*w:styleId="{re.escape(style_id)}"[^>]*>.*?</w:style>',
            styles_xml,
            flags=re.DOTALL,
        )
        hyperlink_styles += len(style_blocks)
        for style_xml in style_blocks:
            colors = re.findall(r'<w:color\b[^>]*w:val="([^"]+)"[^>]*/>', style_xml)
            if not colors or any(color.upper() != BLACK for color in colors) or "w:themeColor=" in style_xml:
                raise ValueError(f"Hyperlink style {style_id} must render in explicit black.")
    if hyperlink_styles < 1:
        raise ValueError("DOCX must define a black hyperlink style.")
    settings = ElementTree.fromstring(settings_xml)
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None or update_fields.get(qn("w:val")) not in {"1", "true"}:
        raise ValueError("DOCX must request field refresh on open.")


def set_cached_field_result(xml: str, instruction: str, value: str) -> tuple[str, int]:
    pattern = re.compile(
        rf"(<w:instrText(?:\s[^>]*)?>\s*{re.escape(instruction)}\s*</w:instrText>"
        rf".*?<w:fldChar[^>]*w:fldCharType=\"separate\"[^>]*/>"
        rf".*?<w:t(?:\s[^>]*)?>)(.*?)(</w:t>)",
        flags=re.DOTALL,
    )
    return pattern.subn(lambda match: f"{match.group(1)}{value}{match.group(3)}", xml)


def set_hyperlink_styles_black(xml: str) -> tuple[str, int]:
    updated_styles = 0
    for style_id in ("Hyperlink", "FollowedHyperlink", "Hyperlink1", "FollowedHyperlink1"):
        style_pattern = re.compile(
            rf'<w:style\b[^>]*w:styleId="{re.escape(style_id)}"[^>]*>.*?</w:style>',
            flags=re.DOTALL,
        )

        def update_style(match: re.Match[str]) -> str:
            nonlocal updated_styles
            style_xml = match.group(0)
            style_xml, color_updates = re.subn(
                r"<w:color\b[^>]*/>",
                '<w:color w:val="000000"/>',
                style_xml,
            )
            if color_updates == 0:
                if "<w:rPr>" in style_xml:
                    style_xml = style_xml.replace("<w:rPr>", '<w:rPr><w:color w:val="000000"/>', 1)
                else:
                    style_xml = style_xml.replace(
                        "</w:style>",
                        '<w:rPr><w:color w:val="000000"/></w:rPr></w:style>',
                    )
            updated_styles += 1
            return style_xml

        xml = style_pattern.sub(update_style, xml)
    return xml, updated_styles


def normalize_hyperlink_styles(path: Path) -> None:
    replacement = path.with_suffix(".styles.tmp.docx")
    updated_styles = 0
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(replacement, "w") as target:
        for entry in source.infolist():
            data = source.read(entry.filename)
            if entry.filename == "word/styles.xml":
                xml, updated_styles = set_hyperlink_styles_black(data.decode("utf-8"))
                data = xml.encode("utf-8")
            target.writestr(entry, data)
    if updated_styles < 1:
        replacement.unlink(missing_ok=True)
        raise ValueError("Could not locate a hyperlink style to normalize.")
    replacement.replace(path)


def finalize_field_cache(path: Path, page_count: int) -> None:
    if page_count <= 0:
        raise ValueError("Final page count must be positive.")
    replacement = path.with_suffix(".fields.tmp.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(replacement, "w") as target:
        updated_num_pages = 0
        for entry in source.infolist():
            data = source.read(entry.filename)
            if entry.filename == "word/settings.xml":
                xml = data.decode("utf-8")
                if re.search(r"<w:updateFields\b", xml):
                    xml = re.sub(
                        r"<w:updateFields\b[^>]*/>",
                        '<w:updateFields w:val="true"/>',
                        xml,
                        count=1,
                    )
                else:
                    xml = xml.replace(
                        "</w:settings>",
                        '<w:updateFields w:val="true"/></w:settings>',
                    )
                data = xml.encode("utf-8")
            elif entry.filename.startswith("word/footer") and entry.filename.endswith(".xml"):
                xml, updates = set_cached_field_result(data.decode("utf-8"), "NUMPAGES", str(page_count))
                updated_num_pages += updates
                data = xml.encode("utf-8")
            target.writestr(entry, data)
    if updated_num_pages < 1:
        replacement.unlink(missing_ok=True)
        raise ValueError("Could not locate a cached NUMPAGES field result.")
    replacement.replace(path)
    normalize_hyperlink_styles(path)


def build(output: Path) -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    validate_source(markdown)
    build_diagrams()
    document = Document()
    configure_document(document)
    properties = document.core_properties
    properties.title = "[DIS] Obscell - Reusable Privacy Infrastructure for CKB Applications"
    properties.subject = "CKB Community DAO application / project development proposal"
    properties.author = "Williams Oluwagbemi Akinwamide (willy264)"
    properties.keywords = "CKB, Nervos, CCC, privacy, SDK, Pudge, grant"
    properties.comments = "Generated from docs/grant-proposal.md. Current settlement status remains explicitly incomplete."
    properties.created = datetime.now(timezone.utc)
    properties.modified = properties.created
    add_cover(document)
    render_markdown(document, markdown)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    normalize_hyperlink_styles(output)
    validate_docx(output)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=OUTPUT)
    parser.add_argument("--validate-only", action="store_true")
    parser.add_argument(
        "--finalize-page-count",
        type=int,
        help="Cache a Word-verified NUMPAGES value and restore update-on-open after Word saves the DOCX.",
    )
    args = parser.parse_args()
    output = args.output.resolve()
    if args.finalize_page_count is not None:
        finalize_field_cache(output, args.finalize_page_count)
        validate_docx(output)
    elif args.validate_only:
        validate_docx(output)
    else:
        build(output)
    print(f"Validated proposal: {output}")


if __name__ == "__main__":
    main()
